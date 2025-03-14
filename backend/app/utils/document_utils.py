"""
Document Utilities
=================

Utilities for extracting text from various document formats.
Supports: PDF, DOCX, TXT, and other common document formats.
"""

import os
import logging
from pathlib import Path
from typing import Optional, Union, List

logger = logging.getLogger(__name__)

def extract_text_from_document(document_path: Union[str, Path]) -> Optional[str]:
    """
    Extract text from a document based on its file extension.
    
    Args:
        document_path: Path to the document file
        
    Returns:
        Extracted text or None if extraction failed
    """
    document_path = Path(document_path)
    
    if not document_path.exists():
        logger.error(f"Document does not exist: {document_path}")
        return None
    
    file_extension = document_path.suffix.lower()
    
    try:
        if file_extension == '.pdf':
            return extract_text_from_pdf(document_path)
        elif file_extension == '.docx':
            return extract_text_from_docx(document_path)
        elif file_extension == '.txt':
            return extract_text_from_txt(document_path)
        elif file_extension in ['.md', '.markdown']:
            return extract_text_from_txt(document_path)
        elif file_extension in ['.html', '.htm']:
            return extract_text_from_html(document_path)
        elif file_extension in ['.csv', '.tsv']:
            return extract_text_from_csv(document_path)
        elif file_extension in ['.json']:
            return extract_text_from_json(document_path)
        else:
            logger.warning(f"Unsupported file format: {file_extension}")
            # Try to extract as plain text as fallback
            return extract_text_from_txt(document_path)
    except Exception as e:
        logger.error(f"Error extracting text from {document_path}: {str(e)}")
        return None

def extract_text_from_pdf(pdf_path: Union[str, Path]) -> str:
    """
    Extract text from a PDF document.
    
    Args:
        pdf_path: Path to the PDF file
        
    Returns:
        Extracted text
    """
    try:
        import PyPDF2
        
        text = []
        with open(pdf_path, 'rb') as file:
            reader = PyPDF2.PdfReader(file)
            for page_num in range(len(reader.pages)):
                page = reader.pages[page_num]
                text.append(page.extract_text())
        
        return "\n\n".join(text)
    except ImportError:
        logger.warning("PyPDF2 not installed. Trying alternative PDF extraction.")
        try:
            # Fallback to pdfplumber if PyPDF2 fails
            import pdfplumber
            
            text = []
            with pdfplumber.open(pdf_path) as pdf:
                for page in pdf.pages:
                    text.append(page.extract_text() or "")
            
            return "\n\n".join(text)
        except ImportError:
            logger.error("PDF extraction libraries (PyPDF2, pdfplumber) not installed.")
            raise ImportError("PDF extraction requires PyPDF2 or pdfplumber. Please install with: pip install PyPDF2 pdfplumber")

def extract_text_from_docx(docx_path: Union[str, Path]) -> str:
    """
    Extract text from a DOCX document.
    
    Args:
        docx_path: Path to the DOCX file
        
    Returns:
        Extracted text
    """
    try:
        import docx
        
        doc = docx.Document(docx_path)
        full_text = []
        
        # Extract text from paragraphs
        for para in doc.paragraphs:
            full_text.append(para.text)
        
        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    row_text.append(cell.text)
                full_text.append(" | ".join(row_text))
        
        return "\n".join(full_text)
    except ImportError:
        logger.error("Python-docx not installed.")
        raise ImportError("DOCX extraction requires python-docx. Please install with: pip install python-docx")

def extract_text_from_txt(txt_path: Union[str, Path]) -> str:
    """
    Extract text from a plain text file.
    
    Args:
        txt_path: Path to the text file
        
    Returns:
        Extracted text
    """
    try:
        encodings_to_try = ['utf-8', 'latin-1', 'cp1252']
        
        for encoding in encodings_to_try:
            try:
                with open(txt_path, 'r', encoding=encoding) as file:
                    return file.read()
            except UnicodeDecodeError:
                continue
        
        # If all encodings fail, use binary mode and decode with errors='replace'
        with open(txt_path, 'rb') as file:
            return file.read().decode('utf-8', errors='replace')
            
    except Exception as e:
        logger.error(f"Error reading text file: {str(e)}")
        raise

def extract_text_from_html(html_path: Union[str, Path]) -> str:
    """
    Extract text from an HTML file.
    
    Args:
        html_path: Path to the HTML file
        
    Returns:
        Extracted text
    """
    try:
        from bs4 import BeautifulSoup
        
        with open(html_path, 'r', encoding='utf-8', errors='replace') as file:
            soup = BeautifulSoup(file.read(), 'html.parser')
            
            # Remove script and style elements
            for script in soup(["script", "style"]):
                script.extract()
            
            # Extract text
            text = soup.get_text(separator="\n")
            
            # Remove excessive newlines
            lines = [line.strip() for line in text.split("\n")]
            text = "\n".join(line for line in lines if line)
            
            return text
    except ImportError:
        logger.warning("BeautifulSoup not installed. Extracting HTML as plain text.")
        return extract_text_from_txt(html_path)

def extract_text_from_csv(csv_path: Union[str, Path]) -> str:
    """
    Extract text from a CSV file.
    
    Args:
        csv_path: Path to the CSV file
        
    Returns:
        Extracted text representation of the CSV
    """
    try:
        import csv
        
        rows = []
        with open(csv_path, 'r', encoding='utf-8', errors='replace') as file:
            csv_reader = csv.reader(file)
            for row in csv_reader:
                rows.append(" | ".join(row))
        
        return "\n".join(rows)
    except Exception as e:
        logger.error(f"Error reading CSV file: {str(e)}")
        return extract_text_from_txt(csv_path)

def extract_text_from_json(json_path: Union[str, Path]) -> str:
    """
    Extract text from a JSON file.
    
    Args:
        json_path: Path to the JSON file
        
    Returns:
        Formatted text representation of the JSON
    """
    try:
        import json
        
        with open(json_path, 'r', encoding='utf-8', errors='replace') as file:
            data = json.load(file)
            # Pretty print the JSON with indentation
            return json.dumps(data, indent=2)
    except Exception as e:
        logger.error(f"Error reading JSON file: {str(e)}")
        return extract_text_from_txt(json_path)

def split_document(text: str, chunk_size: int = 5000, overlap: int = 500) -> List[str]:
    """
    Split a document into overlapping chunks for processing.
    
    Args:
        text: Full document text
        chunk_size: Maximum size of each chunk
        overlap: Number of characters to overlap between chunks
        
    Returns:
        List of text chunks
    """
    if not text:
        return []
    
    # Use sentence boundaries for cleaner splits
    import nltk
    try:
        nltk.data.find('tokenizers/punkt')
    except LookupError:
        nltk.download('punkt')
    
    sentences = nltk.sent_tokenize(text)
    
    chunks = []
    current_chunk = []
    current_length = 0
    
    for sentence in sentences:
        sentence_length = len(sentence)
        
        # If adding this sentence would exceed chunk size, finalize the current chunk
        if current_length + sentence_length > chunk_size and current_chunk:
            chunks.append(" ".join(current_chunk))
            
            # Keep some sentences for overlap
            overlap_sentences = []
            overlap_length = 0
            for sent in reversed(current_chunk):
                if overlap_length + len(sent) <= overlap:
                    overlap_sentences.insert(0, sent)
                    overlap_length += len(sent) + 1  # +1 for space
                else:
                    break
            
            current_chunk = overlap_sentences
            current_length = overlap_length
        
        current_chunk.append(sentence)
        current_length += sentence_length + 1  # +1 for space
    
    # Add the last chunk if it's not empty
    if current_chunk:
        chunks.append(" ".join(current_chunk))
    
    return chunks 